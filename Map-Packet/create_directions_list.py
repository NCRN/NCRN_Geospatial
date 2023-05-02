#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Abstract: Initial working script for creating directions to veg monitoring locations.

Description:
The purpose of this script is to create directions sheets for NCRN Vegetation Monitoring.
An excel spreadsheet will define the parameters for populating a Word template.
Each output docx contains monitoring locations with the same "Map" and "Area" for each panel (1-4)
Finally, the script will convert the Word docx to PDFs so that they can be included in the map packet.

TODO: Additional Refactoring
--------------------------------------------------------------------------------
References:
# https://betterprogramming.pub/automate-filling-templates-with-python-1ff6c6fd595e

#__author__ = "NCRN GIS Unit"
#__copyright__ = "None"
#__credits__ = [""]
#__license__ = "GPL"
#__version__ = "1.0.2"
#__maintainer__ = "David Jones"
#__email__ = "david_jones@nps.gov"
#__status__ = "Staging"
"""

# Import statements for utilized libraries / packages
from cmath import nan
from logging import exception
from tkinter.tix import ROW
from openpyxl import load_workbook
from docx import Document
from docx.enum.section import WD_ORIENT
from docxtpl import DocxTemplate
from docx.shared import RGBColor
import pandas as pd
import sys
from reportlab.pdfgen.canvas import Canvas
import comtypes.client
import time
import fitz
import shutil
import os
from os import path

print("### GETTING STARTED ###".format())

"""
Set various global variables. Some of these could be parameterized to be used in an 
ArcGIS Toolbox script and/or command line use. 
"""

# Currently hardcoded values that may be parameterized if bundling into a tool
__ROOT_DIR = r'C:\Users\goettel\DOI\NPS-NCRN-Forest Veg - Documents\General\Field_Maps_2023' ## Set the directory path to the root directory that will be the destination for outputs. NEED TO UPDATE PREFIX TO YOUR ONEDRIVE ACCOUNT
template_path = r'C:\Users\goettel\DOI\NCRN Data Management - Geospatial\Templates\NCRN\Directions_Template.docx' ## Create a variable to store the full path to the Word doc template. NEED TO UPDATE PREFIX TO YOUR ONEDRIVE ACCOUNT
__XCEL_LIBRARY = r'C:\Users\goettel\DOI\NCRN Data Management - Geospatial\NCRN_Veg_Locations_Directions_Master.xlsx' ## Create a variable to store the full path to the Excel file. NEED TO UPDATE PREFIX TO YOUR ONEDRIVE ACCOUNT

# Create some functions to be used in various places
def Clear_Folder(folder_path):
    """Takes the path of a folder or geodatabase and deletes the files within it"""
    for filename in os.listdir(folder_path):
        fullpath_filename = os.path.join(folder_path, filename)
        if os.path.isfile(fullpath_filename):
            try:
                os.remove(fullpath_filename)
            except:
                print("ERROR CLEARING DIRECTIONS FOLDERS! Word may be running in the background and using the files. RESTART COMPUTER AND TRY AGAIN...")

def create_blank_docx(panel_folder, Map_List):
    """
    Use to create a blank Word .docx
    e.g. panel_folder = 'Panel_1' (The panel folder the maps are located in. All maps must be in the same panel folder.)
    e.g. Map_List = 'Map_Name1', 'Map_Name3', 'Map_Name5' (The maps that a blank document will be placed behind)
    """
    for Map in Map_List:
        document = Document()
        # Change orientation of the output document to landscape
        sections = document.sections
        section = sections[0]
        new_width, new_height = section.page_height, section.page_width
        section.page_width = new_width
        section.page_height = new_height
        section.orientation = WD_ORIENT.LANDSCAPE
        # '03' orders blank page after the directions sheets
        filename = Map + " 03 " + "blank page.docx"
        fullpath_filename = os.path.join(__ROOT_DIR, panel_folder, 'Directions', filename)
        document.save(fullpath_filename)

def covx_to_pdf(in_file, out_file):
    """
    Convert a Word .docx to PDF
    Requires the name of the input Word .docx and the name of the output PDF
    """
    wdFormatPDF = 17
    word = comtypes.client.CreateObject('Word.Application')
    # Keep the program hidden while running in the background
    word.Visible = False
    doc = word.Documents.Open(in_file)
    # add a 7 second delay so that the program doesn't crash
    time.sleep(7)
    doc.SaveAs(out_file, FileFormat = wdFormatPDF)
    doc.Close()
    word.Quit()

def pdf_delete_page(panel_folder, filename, page_to_delete):
    ''' 
    Deletes a single page from a PDF
    e.g. page_to_delete = 1
    The page associated with the number is deleted in the specified PDF. indexing starts from 1 so if we pass 1 as an argument the second page will be deleted.
    '''
    folder_path = os.path.join(__ROOT_DIR, panel_folder, 'Directions')
    input_file = os.path.join(folder_path, filename)
    root_ext = os.path.splitext(filename)
    root = root_ext[0]
    new_filename = root + " update" + '.pdf'
    output_file = os.path.join(folder_path, new_filename)
    # try/except will skip if the page has already been deleted
    try:
        file_handle = fitz.open(input_file)
        file_handle.delete_page(page_to_delete)
        file_handle.save(output_file)
        file_handle.close()
        os.remove(input_file)
    except:
        pass

# Set the imported docx as the template
template = DocxTemplate(template_path)

# Create dataframes for Areas and Plots using the imported spreadsheet
df_areas = pd.read_excel(__XCEL_LIBRARY, sheet_name='Areas Python')
df_plots = pd.read_excel(__XCEL_LIBRARY, sheet_name='Plots', na_filter = False)

# Create a list with the four panel folders to loop over
panel_list = 'Panel_1', 'Panel_2', 'Panel_3', 'Panel_4'

# Remove all old files from the directions folder
for panel in panel_list:
    Clear_Folder(os.path.join(__ROOT_DIR, panel, 'Directions'))

# Create empty list to populate with plots
plots_populated_list = []

# Loop over rows in the Areas dataframe to create directions Word docx
print("Creating Directions docx...")
for index, row in df_areas.iterrows():
    Area = row['Area']
    Panel = row['Panel']
    Map = row['Map']
    # Display '0' cells as blank
    if row['Area Warnings'] == 0:
        Area_Warnings = ""
    else: 
        Area_Warnings = row['Area Warnings']
    # This is a dictionary with the first 3 keys matching columns in the Area dataframe. This will fill out the Area portion of the template.
    to_fill_in_area = {
        'Map': Map,
        'Area': Area,
        'Area_Directions': row['Area Directions'],
        'Area_Warnings': Area_Warnings,
        # The next keys refill the Plots input parameters (they are erased by the first render function)
        'Plot': '{{Plot_Name}}',
        'Plot_Directions': '{{Plot_Directions}}',
        'Warnings': '{{Warnings}}',
        'Keys': '{{Keys}}',
        }
    # Use the render function to fill in the word template based on the dictionary created.
    template.render(to_fill_in_area)
    # '02' orders directions sheets after the map
    filename = Map + " 02 " + Area + '.docx'
    filled_path = os.path.join(__ROOT_DIR, row['Panel_Folder'], 'Directions', filename)
    # save the created template on the filled path
    try:
        template.save(filled_path)
        # Loop over rows in the Plots dataframe to populate the docx
        for index, row in df_plots.iterrows():
            # Change the template to the docx just created
            filled_template = DocxTemplate(filled_path)
            # Only populate docx if the area, panel, and map matches
            if (row['Area'] == Area) & (row['Panel'] == Panel) & (row['Map'] == Map):
                # This is a dictionary with the keys matching columns in the Plots dataframe
                to_fill_in_plot = {
                    'Plot_Name': row['Plot Name'],
                    'Plot_Directions': row['Plot Directions'],
                    'Warnings': row['Plot Warnings'],
                    'Keys': row['Keys'],
                    }
                plots_populated_list.append(row['Plot Name'])
                filled_template.render(to_fill_in_plot)
                # Create a new row in the Plots table to be filled out in the next itteration
                table = filled_template.tables[0]
                row = table.add_row().cells
                row[0].text = '{{Plot_Name}}'
                row[1].text = '{{Plot_Directions}}'
                row[2].text = '{{Warnings}}'
                # Add red font to plot warnings
                warnings_paragraph = row[2].paragraphs[0]
                warnings_run = warnings_paragraph.runs
                warnings_font = warnings_run[0].font
                warnings_font.color.rgb = RGBColor.from_string('FF0000')
                row[3].text = '{{Keys}}'
                filled_template.save(filled_path)
            # Delete input parameters that were left empty
            to_fill_in = {}
        filled_template.render(to_fill_in)
        filled_template.save(filled_path)
    except:
        print("ERROR CREATING '{0}' in {1}".format(filename, panel))
    
# Identify plots that weren't added to a directions sheet
plots_expected_df = df_plots['Plot Name']
plots_expected_list = plots_expected_df.values.tolist()

set_difference = set(plots_expected_list) - set(plots_populated_list)
list_difference_result = list(set_difference)
print("Plots not added to Directions: ", list_difference_result)

# Create blank docx to go behind map
# Need to run for maps that have an even number of directions sheets (2, 4, etc.)
# Panel 1
Blank_Pages_1 = "CHOH 4 (Point of Rocks)", "CHOH 5 and HAFE", "GWMP (Mt Vernon) and NACE South (FOWA, PISC)", "MONO", "NACE (ANAC, FODU)"
create_blank_docx('Panel_1', Blank_Pages_1)

# Panel 2
Blank_Pages_2 = "CATO", "CHOH 2 and GWMP (GRFA)", "CHOH 5 and HAFE", "CHOH 9 (Far West - Paw Paw Tunnel)", "GWMP (Daingerfield) and NACE (Shepherd Parkway)", "MONO", "NACE (FODU, SUIT)"
create_blank_docx('Panel_2', Blank_Pages_2)

# Panel 3
Blank_Pages_3 = "CATO", "CHOH 1, 2 and GWMP (GRFA, Turkey Run)", "CHOH 10 (Far West - Oldtown, Spring Gap)", "CHOH 5 and HAFE", "GWMP (Mt Vernon) and NACE South (PISC)", "MONO", "NACE (BAWA, GREE)", "PRWI"
create_blank_docx('Panel_3', Blank_Pages_3)

# Panel 4
Blank_Pages_4 = "ANTI and CHOH 6 (Snyders Landing)", "CHOH 2 and GWMP (Turkey Run)", "CHOH 3 (Violettes Lock) and GWMP (GRFA)", "CHOH 5 and HAFE"
create_blank_docx('Panel_4', Blank_Pages_4)

# Loop over the panel folders to convert docx to pdf
for panel in panel_list:
    print("Converting {0} Directions docx to pdf...".format(panel))
    folder_path = os.path.join(__ROOT_DIR, panel, 'Directions')
    for filename in os.listdir(folder_path):
        if filename.endswith('.docx'):
            # Full path of the docx
            in_file = os.path.join(folder_path, filename)
            # Name of the output pdf
            out_filename = filename.replace('.docx', '.pdf')
            # Full path of the output pdf
            out_file = os.path.join(folder_path, out_filename)
            # Create an empty PDF
            canvas = Canvas(out_file)
            # Populate the pdf with the docx
            try:
                covx_to_pdf(in_file, out_file)
            except:
                print("ERROR CONVERTING {0} TO PDF".format(in_file))

# Remove empty pages at the end of PDFs
pdf_delete_page('Panel_1', "PRWI 02 PRWI.pdf", 1)



