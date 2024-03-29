#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Abstract: Initial working script to create map packets for vegetation monitoring

Description:
1) The script will create directions and appendix sheets for NCRN Vegetation Monitoring.
An excel spreadsheet will define the parameters for populating the two Word templates.
The script will convert the docx to PDF so that they can be included in the map packet.

2) The script will create the map/directions packet for NCRN vegetation monitoring.
The script will find the materials that make up the map packet in their respective folders and combine them

TODO: Additional Refactoring
--------------------------------------------------------------------------------
References:
# https://python-bloggers.com/2022/04/merging-pdfs-with-python-2/
# https://betterprogramming.pub/automate-filling-templates-with-python-1ff6c6fd595e

#__author__ = "NCRN GIS Unit"
#__copyright__ = "None"
#__credits__ = [""]
#__license__ = "GPL"
#__version__ = "1.0.2"
#__maintainer__ = "David Jones"
#__email__ = "david_jones@nps.gov"
#__status__ = "Staging"
"""

# Import statements for utilized libraries / packages
# Script is still in draft mode, some of these packages may not have been used
from sqlite3 import paramstyle
import PyPDF2
from PyPDF2 import PdfFileReader, utils
from io import StringIO
import subprocess
import os
import shutil
from cmath import nan
from logging import exception
from tkinter.tix import ROW
from openpyxl import load_workbook
import docx
from docx import Document
from docx.shared import Pt
from docx.enum.section import WD_ORIENT
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docxtpl import DocxTemplate
from docx.shared import RGBColor
import pandas as pd
import sys
from reportlab.pdfgen.canvas import Canvas
import comtypes.client
import time
import fitz
from os import path
import win32com.client

print("### GETTING STARTED ###".format())

"""
Set various global variables. Some of these could be parameterized to be used in an 
ArcGIS Toolbox script and/or command line use. 
"""
# Currently hardcoded values that may be parameterized if bundling into a tool
__ROOT_DIR = r'C:\Users\goettel\DOI\NPS-NCRN-Forest Veg - Documents\General\Field_Maps_2023' ## Set the directory path to the root directory that will be the destination for outputs. NEED TO UPDATE PREFIX TO YOUR ONEDRIVE ACCOUNT
__XCEL_LIBRARY = r'C:\Users\goettel\OneDrive - DOI\Geospatial\NCRN_Veg_Locations_Directions_Master.xlsx' ## Create a variable to store the full path to the Excel file. NEED TO UPDATE PREFIX TO YOUR ONEDRIVE ACCOUNT
directions_template_path = r'C:\Users\goettel\DOI\NCRN Data Management - Geospatial\Templates\NCRN\Veg_Map_Packet\Directions_Template.docx' ## Create a variable to store the full path to the Directions template. NEED TO UPDATE PREFIX TO YOUR ONEDRIVE ACCOUNT
appendix_template_path = r'C:\Users\goettel\DOI\NCRN Data Management - Geospatial\Templates\NCRN\Veg_Map_Packet\Appendix_Template.docx' ## Create a variable to store the full path to the Appendix template. NEED TO UPDATE PREFIX TO YOUR ONEDRIVE ACCOUNT

# Create some functions to be used in various places
def Clear_Folder(folder_path):
    """Takes the path of a folder or geodatabase and deletes the files within it"""
    for filename in os.listdir(folder_path):
        fullpath_filename = os.path.join(folder_path, filename)
        if os.path.isfile(fullpath_filename):
            try:
                os.remove(fullpath_filename)
            except:
                print("ERROR CLEARING FOLDERS! Word or Adobe may be running in the background and using the files. RESTART COMPUTER AND TRY AGAIN...")

def create_blank_docx(folder_path, map_list):
    """
    Use to create a blank Word .docx
    Requires: 
    1) The folder the maps are located in.
    2) A list of maps to add a blank page behind. e.g. Map_List = 'Map_Name1', 'Map_Name3', 'Map_Name5'
    """
    for Map in map_list:
        document = Document()
        # Change orientation of the output document to landscape
        sections = document.sections
        section = sections[0]
        new_width, new_height = section.page_height, section.page_width
        section.page_width = new_width
        section.page_height = new_height
        section.orientation = WD_ORIENT.LANDSCAPE
        # '03' orders blank page after the directions sheets
        filename = Map + " 03 " + "blank page.docx"
        fullpath_filename = os.path.join(folder_path, filename)
        document.save(fullpath_filename)

#def covx_to_pdf(input_fullpath_filename, output_fullpath_filename):
#    """
#    Convert a Word (.docx) to PDF (.pdf)
#    Requires the full file path of the input .docx and the output PDF
#    """
#    wdFormatPDF = 17
#    word = comtypes.client.CreateObject('Word.Application')
#    # Keep the program hidden while running in the background
#    word.Visible = False
#    doc = word.Documents.Open(input_fullpath_filename)
#    # add a 7 second delay so that the program doesn't crash
#    time.sleep(7)
#    doc.SaveAs(output_fullpath_filename, FileFormat = wdFormatPDF)
#    doc.Close()
#    word.Quit()

def covx_to_pdf(input_fullpath_filename, output_fullpath_filename):
    """
    Convert a Word (.docx) to PDF (.pdf)
    Requires the full file path of the input .docx and the output PDF
    """
    wdFormatPDF = 17
    word = win32com.client.Dispatch('Word.Application')
    # Keep the program hidden while running in the background
    word.Visible = False
    doc = word.Documents.Open(input_fullpath_filename)
    # add a 15 second delay so that the program doesn't crash
    time.sleep(15)
    doc.SaveAs(output_fullpath_filename, FileFormat = wdFormatPDF)
    doc.Close()
    word.Quit()

def pdf_delete_page(folder_path, filename, page_to_delete):
    ''' 
    Deletes a single page from a PDF
    e.g. page_to_delete = 1
    The page associated with the number is deleted from the PDF. indexing starts from 1 so if we pass 1 as an argument the second page will be deleted.
    '''
    input_file = os.path.join(folder_path, filename)
    root_ext = os.path.splitext(filename)
    root = root_ext[0]
    new_filename = root + " update" + '.pdf'
    output_file = os.path.join(folder_path, new_filename)
    file_handle = fitz.open(input_file)
    file_handle.delete_page(page_to_delete)
    file_handle.save(output_file)
    file_handle.close()
    os.remove(input_file)

def add_hyperlink(paragraph, url, text):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )
    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')
    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run ()
    r._r.append (hyperlink)
    # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True
    return hyperlink

def add_map_suffix(map_folder_path):
    """
    Adds '01' to the end of all PDF files in the given map folder 
    '01' orders map before the directions sheets
    """
    for filename in os.listdir(map_folder_path):
        # Skip if already has '01'
        if filename.endswith('01.pdf'):
            pass
        else:
            if filename.endswith('.pdf'):
                src = os.path.join(map_folder_path, filename)
                root_ext = os.path.splitext(filename)
                root = root_ext[0]
                new_filename = root + ' 01' + '.pdf'
                dst = os.path.join(map_folder_path, new_filename)
                os.rename(src, dst)

def create_pdf_list(pdf_folder_path, pdf_list_name):
    """
    Creates a list of PDFs
    Requires (1) the folder containg the pdfs and (2) what you want to call the output list
    """
    for filename in os.listdir(pdf_folder_path):
        # only PDFs are added to the list
        if filename.endswith('.pdf'):
            pdf_list_name.append(os.path.join(pdf_folder_path, filename))
    # Sorts list ABC
    pdf_list_name.sort(key = str.lower)

def merge_pdfs(pdf_list, pdf_filename):
    """
    Takes a list of PDFs and merges them into one PDF
    Requires (1) a list of PDFs (2) what you want to name the output PDF
    """
    pdfMerge = PyPDF2.PdfFileMerger()
    for pdf in pdf_list:
        pdfFile = open(pdf, 'rb')
        pdfReader = PyPDF2.PdfFileReader(pdfFile)
        pdfMerge.append(pdfReader)
    pdfFile.close()
    pdfMerge.write(os.path.join(__ROOT_DIR, 'Packets', pdf_filename))

def copy_paste_pdfs(source_folder, destination_folder):
    """
    Copies all PDFs from one folder to another
    """
    for filename in os.listdir(source_folder):
        source = os.path.join(source_folder, filename)
        destination = os.path.join(destination_folder, filename)
        # only PDFs are copied
        if filename.endswith('.pdf'):
            shutil.copy(source, destination)    

def pdf_port_to_land(portrait_folder, landscape_folder):
    # takes a folder of PDFs with portrait orientation
    # rotates all PDFs to landscape orientation
    # exports to a new folder
    for filename in os.listdir(portrait_folder):
        if filename.endswith('.pdf'):
            input_fullpath_filename = os.path.join(portrait_folder, filename)
            output_fullpath_filename = os.path.join(landscape_folder, filename)
            pdf_in = open(input_fullpath_filename, 'rb')
            pdf_reader = PyPDF2.PdfFileReader(pdf_in)
            pdf_writer = PyPDF2.PdfFileWriter()
            numofpages = pdf_reader.numPages
            numrotated = 0 
            for pagenum in range(numofpages):
                page = pdf_reader.getPage(pagenum)
                mb = page.mediaBox
                if (mb.upperRight[1] > mb.upperRight[0]) and (page.get('/Rotate') is None):
                    page.rotateCounterClockwise(90)
                    numrotated = numrotated + 1
                pdf_writer.addPage(page)
            pdf_out = open(output_fullpath_filename, 'wb')
            pdf_writer.write(pdf_out)
            pdf_out.close()
            pdf_in.close()

# Create a list with the four panel folders to loop over
panel_list = 'Panel_1', 'Panel_2', 'Panel_3', 'Panel_4'

# Create dataframes & lists for Areas and Plots using the imported spreadsheet
df_areas = pd.read_excel(__XCEL_LIBRARY, sheet_name='Areas Python')
df_plots = pd.read_excel(__XCEL_LIBRARY, sheet_name='Plots', na_filter = False)
plots_expected_df = df_plots['Plot Name']
plots_expected_list = plots_expected_df.values.tolist()

# Create empty list to populate with plots not added to Directions
plots_populated_list = []

# Create empty list to populate with plots not added to Appendix
plots_populated_appendix_list = []

# Prepare folders and move files to the correct locations
for panel in panel_list:
    ## Clear old files
    ## Directions pdfs
    Clear_Folder(os.path.join(__ROOT_DIR, panel, 'Directions'))
    # Directions word docx
    Clear_Folder(os.path.join(__ROOT_DIR, panel, 'Directions', 'Working'))
    ## maps working
    folder_maps_working = os.path.join(__ROOT_DIR, panel, 'Maps', 'Working')
    Clear_Folder(folder_maps_working)
    ## maps
    folder_maps = os.path.join(__ROOT_DIR, panel, 'Maps')
    Clear_Folder(folder_maps)
    ## maps/directions
    folder_maps_directions = os.path.join(__ROOT_DIR, panel, 'Working')
    Clear_Folder(folder_maps_directions)
    ## Appendix
    Clear_Folder(os.path.join(__ROOT_DIR, panel, 'Appendix'))
    Clear_Folder(os.path.join(__ROOT_DIR, panel, 'Appendix', 'Working'))

    # copy maps to main Maps folder. This is where they will be easy to view for in-field purposes.
    folder_maps_port = os.path.join(folder_maps_working, 'Portrait')
    folder_maps_land = os.path.join(folder_maps_working, 'Landscape')
    copy_paste_pdfs(folder_maps_port, folder_maps)
    copy_paste_pdfs(folder_maps_land, folder_maps)
    # copy landscape map pdfs to maps working folder
    copy_paste_pdfs(folder_maps_land, folder_maps_working)
    # rotate portrait pdf maps to landscape and copy to maps working folder
    pdf_port_to_land(folder_maps_port, folder_maps_working)
    # add '01' suffix to pdfs in the Maps folder
    add_map_suffix(folder_maps_working)
    ## Copy map pdfs to the maps/directions folder
    copy_paste_pdfs(folder_maps_working, folder_maps_directions)

# Loop over rows in the Areas dataframe to create directions Word docx
print("Creating Directions docx...")
for index, row in df_areas.iterrows():
    # Set the Directions docx as the template
    area_template = DocxTemplate(directions_template_path)
    Area = row["Area"]
    Panel = row["Panel Folder"]
    Map = row["Map"]
    # Display '0' cells as blank
    if row["Area Warnings"] == 0:
        Area_Warnings = ""
    else:
        Area_Warnings = row["Area Warnings"]
    # This is a dictionary with the first 4 keys referencing columns in the Area dataframe. This will fill out the header portion of the template.
    to_fill_in_area = {
        'Map': Map,
        'Area': Area,
        'Area_Directions': row["Area Directions"],
        'Area_Warnings': Area_Warnings,
        # The next keys regenerate the Plots input parameters (they are erased by the first render function)
        'Plot': '{{Plot_Name}}',
        'Plot_Directions': '{{Plot_Directions}}',
        'Warnings': '{{Warnings}}',
        'Keys': '{{Keys}}'
        }
    # Use the render function to fill in the template based on the dictionary created.
    area_template.render(to_fill_in_area)
    # '02' orders directions sheets after the map
    filename = Map + " 02 " + Area + '.docx'
    filled_path = os.path.join(__ROOT_DIR, Panel, 'Directions', 'Working', filename)
    # save the created template on the filled path
    area_template.save(filled_path)
    # Loop over rows in the Plots dataframe to populate the docx
    for index, row in df_plots.iterrows():
        # Change the template to the docx just created
        filled_template = DocxTemplate(filled_path)
        # Only populate docx if the area, panel, and map matches
        if (row['Area'] == Area) & (row['Panel Folder'] == Panel) & (row['Map'] == Map):
            # This is a dictionary with the keys matching columns in the Plots dataframe
            to_fill_in_plot = {
                'Plot_Name': row["Plot Name"],
                'Plot_Directions': row["Plot Directions"],
                'Warnings': row["Plot Warnings"],
                'Keys': row["Keys"]
                }
            plots_populated_list.append(row['Plot Name'])
            filled_template.render(to_fill_in_plot)
            # Create a new row in the Plots table to be filled out in the next itteration
            table = filled_template.tables[0]
            row = table.add_row().cells
            row[0].text = "{{Plot_Name}}"
            row[1].text = "{{Plot_Directions}}"
            row[2].text = "{{Warnings}}"
            row[3].text = "{{Keys}}"
            # Add red font to plot warnings
            warnings_paragraph = row[2].paragraphs[0]
            warnings_run = warnings_paragraph.runs
            warnings_font = warnings_run[0].font
            warnings_font.color.rgb = RGBColor.from_string('FF0000')                
            filled_template.save(filled_path)
    # Delete input parameters that were left empty
    to_fill_in_empty = {}
    filled_template.render(to_fill_in_empty)
    filled_template.save(filled_path)
    
# Identify plots that weren't added to a directions sheet
set_difference_directions = set(plots_expected_list) - set(plots_populated_list)
list_difference_result_directions = list(set_difference_directions)
print("Plots not added to Directions: ", list_difference_result_directions)

for panel in panel_list:
    folder_directions_working = os.path.join(__ROOT_DIR, panel, 'Directions', 'Working')
    folder_directions = os.path.join(__ROOT_DIR, panel, 'Directions')
    
    # Create blank docx to go behind map
    # Need to run for maps that have an even number of directions sheets (2, 4, etc.)
    if panel == 'Panel_1':
        map_list = "CHOH 4 (Point of Rocks)", "CHOH 5 and HAFE", "GWMP (Mt Vernon) and NACE South (FOWA, PISC)", "MONO", "NACE (ANAC, FODU)"
    elif panel == 'Panel_2':
        map_list = "CATO", "CHOH 2 and GWMP (GRFA)", "CHOH 5 and HAFE", "CHOH 9 (Far West - Paw Paw Tunnel)", "GWMP (Daingerfield) and NACE (Shepherd Parkway)", "MONO", "NACE (FODU, SUIT)"
    elif panel == 'Panel_3':
        map_list = "CHOH 1, 2 and GWMP (GRFA, Turkey Run)", "CHOH 10 (Far West - Oldtown, Spring Gap)", "CHOH 5 and HAFE", "GWMP (Mt Vernon) and NACE South (PISC)", "MONO", "NACE (BAWA, GREE)", "PRWI"
    elif panel == 'Panel_4':
        map_list = "ANTI and CHOH 6 (Snyders Landing)", "CHOH 2 and GWMP (Turkey Run)", "CHOH 3 (Violettes Lock) and GWMP (GRFA)", "CHOH 5 and HAFE"
    create_blank_docx(folder_directions_working, map_list)

    # Loop over the panel folders to convert Directions docx to pdf
    print("Converting {0} directions docx to pdf...".format(panel))
    for filename in os.listdir(folder_directions_working):
        # Full path of the docx
        input_fullpath_filename = os.path.join(folder_directions_working, filename)
        # Name of the output pdf
        output_filename = filename.replace('.docx', '.pdf')
        # Full path of the output pdf
        output_fullpath_filename = os.path.join(folder_directions, output_filename)
        # Create an empty PDF
        canvas = Canvas(output_fullpath_filename)
        covx_to_pdf(input_fullpath_filename, output_fullpath_filename)

    # Remove empty pages at the end of PDFs
    for filename in os.listdir(folder_directions):
        if ((panel == 'Panel_1') & (filename == "PRWI 02 PRWI (Park Entrance).pdf")):
            pdf_delete_page(folder_directions, filename, 1)
        elif ((panel == 'Panel_4') & (filename == "CATO 02 CATO.pdf")):
            pdf_delete_page(folder_directions, filename, 1)
    
    # Copy directions pdfs to the Maps/Directions folder
    folder_maps_directions = os.path.join(__ROOT_DIR, panel, 'Working')
    copy_paste_pdfs(folder_directions, folder_maps_directions)

## Create the Appendix
  
# Loop over rows in the Plots dataframe to create Appendix Word docx
print("Creating Appendix docx...")
for panel in panel_list:
    filled_path = os.path.join(__ROOT_DIR, panel, 'Appendix', 'Working', 'Appendix.docx')
    # Copy Appendix docx to Appendix folder for each panel
    shutil.copyfile(appendix_template_path, filled_path)
    # Set the Appendix docx as the template
    appendix_template = DocxTemplate(filled_path)
    for index, row in df_plots.iterrows():
        Panel_Folder = row["Panel Folder"]
        if Panel_Folder == panel:
            plots_populated_appendix_list.append(row['Plot Name'])
            # This is a dictionary with the keys matching columns in the Plots dataframe. This will fill out the table in the Appendix template.                        
            to_fill_in = {
                'Parking_Lat_Long': row["Parking Lat, Long"],
                'Plot_Lat_Long': row["Plot Lat, Long"],
                'Map_Number': row["Map #"]
                }
            # Use the render function to set the appendix template as the working doc
            appendix_template.render(to_fill_in)
            # Identify the table and cells to add
            table = appendix_template.tables[0]
            row_cells = table.add_row().cells
            # Add hyperlink to cells in the first column
            p_table = row_cells[0].paragraphs[0]
            hyperlink = add_hyperlink(p_table, row["Parking Location (Google)"], row["Plot Name"])
            # Create a row in the Plots table to be filled out
            row_cells[1].text = "{{Parking_Lat_Long}}"
            row_cells[2].text = "{{Plot_Lat_Long}}"
            row_cells[3].text = "{{Map_Number}}"
            appendix_template.save(filled_path)
            # use the render function to fill in the word template based on the dictionary created
            appendix_template.render(to_fill_in)
            appendix_template.save(filled_path)

# Find what plots were excluded
set_difference_appendix = set(plots_expected_list) - set(plots_populated_appendix_list)
list_difference_result_appendix = list(set_difference_appendix)
print("Plots not added to Appendix: ", list_difference_result_appendix)

# Loop over the panel folders to convert Appendix docx to pdf
for panel in panel_list:
    print("Converting {0} Appendix docx to pdf...".format(panel))
    folder_appendix_working = os.path.join(__ROOT_DIR, panel, 'Appendix', 'Working')
    folder_appendix = os.path.join(__ROOT_DIR, panel, 'Appendix')
    for filename in os.listdir(folder_appendix_working):
        # Full path of the docx
        input_fullpath_filename = os.path.join(folder_appendix_working, filename)
        # Name of the output pdf
        output_filename = filename.replace('.docx', '.pdf')
        # Full path of the output pdf
        output_fullpath_filename = os.path.join(folder_appendix, output_filename)
        # Create an empty PDF
        canvas = Canvas(output_fullpath_filename)
        # Populate the pdf with the docx
        # Try statement to ignore "Working" folder
        covx_to_pdf(input_fullpath_filename, output_fullpath_filename)
    for filename in os.listdir(folder_appendix):
        # Delete empty pages
        # try statement to ignore 'Working' folder
        try:
            pdf_delete_page(folder_appendix, filename, 3)
        except:
            pass
        # rotate pdfs to landscape and copy back to working folder
        pdf_port_to_land(folder_appendix, folder_appendix_working)

# Create the Map packet

# Map Legend
folder_legend = os.path.join(__ROOT_DIR, 'Map_Legend')
pdfiles_legend = []
create_pdf_list(folder_legend, pdfiles_legend)

for panel in panel_list:
    print("Creating packet for {0}...".format(panel))
    
    # Cover Page
    folder_cover = os.path.join(__ROOT_DIR, panel, 'Cover_Page')
    pdfiles_cover = []
    create_pdf_list(folder_cover, pdfiles_cover)
    
    # Overview map
    folder_overview_map = os.path.join(__ROOT_DIR, panel, 'Overview_Map')
    pdfiles_overview_map = []
    create_pdf_list(folder_overview_map, pdfiles_overview_map)    
    # Overview map blank back
    folder_overview_blank = os.path.join(folder_overview_map, 'Working')
    pdfiles_overview_blank = []
    create_pdf_list(folder_overview_blank, pdfiles_overview_blank)
    # Combine Overview map and blank back
    pdfiles_overview = pdfiles_overview_map + pdfiles_overview_blank
    
    # Maps/directions
    folder_maps_directions = os.path.join(__ROOT_DIR, panel, 'Working')
    pdfiles_maps_directions = [] 
    create_pdf_list(folder_maps_directions, pdfiles_maps_directions)   
    
    ## Appendix
    folder_appendix_working = os.path.join(__ROOT_DIR, panel, 'Appendix', 'Working')
    pdfiles_appendix = []
    create_pdf_list(folder_appendix_working, pdfiles_appendix)

    # Create a giant list with everything to include in the merged PDF
    pdfiles_packet = pdfiles_cover + pdfiles_legend + pdfiles_overview + pdfiles_maps_directions + pdfiles_appendix
    filename = 'Packet_' + panel + '.pdf'
    merge_pdfs(pdfiles_packet, filename)
    print("Done")

                        
